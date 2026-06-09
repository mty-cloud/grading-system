#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
作业批改系统 - Web 应用
支持上传Word文档、提取学生信息和截图、智谱AI视觉评分、导出Excel
打包为独立exe，无需Python环境
"""

import os
import re
import json
import base64
import sys
import webbrowser
import threading
import time
from datetime import datetime

# ===== PyInstaller 打包后的资源路径处理 =====
def resource_path(relative_path):
    """获取资源文件路径（兼容打包后和开发环境）"""
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

# ===== API Key 配置（已内置，开箱即用） =====
ZHIPU_API_KEY = "ae0c33227c2542659ae03872244959cb.hcjuR7slRxiPPROU"
ZHIPU_MODEL = "glm-4v-flash"

# 检查 exe 同目录是否有 .env（仅用于覆盖默认值）
_env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
if os.path.exists(_env_path):
    try:
        with open(_env_path, 'r', encoding='utf-8') as _f:
            for _line in _f:
                _line = _line.strip()
                if _line and not _line.startswith('#') and '=' in _line:
                    _k, _v = _line.split('=', 1)
                    if _k.strip() == 'ZHIPU_API_KEY':
                        ZHIPU_API_KEY = _v.strip()
                    elif _k.strip() == 'ZHIPU_MODEL':
                        ZHIPU_MODEL = _v.strip()
    except Exception:
        pass
# ==============================

import requests
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, jsonify
import docx
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

# Flask 应用
app = Flask(__name__,
            template_folder=resource_path('templates'),
            static_folder=resource_path('static'))
app.secret_key = "grading_secret_key_2024"

# 路径配置（数据/导出文件放在exe同目录，资源模板在打包目录内）
def _exe_dir():
    """获取exe所在目录（打包后）或源码目录（开发时）"""
    try:
        return os.path.dirname(os.path.abspath(sys.executable))
    except Exception:
        return os.path.dirname(os.path.abspath(__file__))

BASE_DIR = _exe_dir()
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
IMAGE_FOLDER = os.path.join(BASE_DIR, "images")
REPORTS_FOLDER = os.path.join(BASE_DIR, "reports")
REPORT_FILE = os.path.join(BASE_DIR, "批改报告.html")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(IMAGE_FOLDER, exist_ok=True)
os.makedirs(REPORTS_FOLDER, exist_ok=True)

# ========== 智谱AI 视觉评分模块 ==========

ZHIPU_BASE_URL = "https://open.bigmodel.cn/api/paas/v4/"

def build_grading_prompt(text):
    prompt = f"""你是一个严格的大学实验报告评分助手。请分析截图和文字内容。

## 学生文字内容
```
{text[:3000]}
```

## 第一步：定等级（必须从下面5个等级中选择1个）
A级 - 优秀(90-95分): 操作完全正确，截图清晰完整，文字充实有深度
B级 - 良好(75-84分): 操作基本正确，截图完整，文字完整无硬伤
C级 - 中等(65-74分): 操作有缺陷或报错，截图不完整，文字有重复
D级 - 及格(50-64分): 明显缺陷，截图缺失或质量差，文字敷衍
E级 - 差(0-49分): 严重缺失，几乎无截图，内容空洞

## 第二步：在等级区间内精确打分。A级需很优秀才能得90+

## 扣分硬规则
- 截图少于2张：降一级
- 文字少于500字：降一级
- 完全没有实验总结：总分最高70
- 内容大量重复：降一级

## 输出要求
纯JSON：
{{"grade":"A/B/C/D/E","image_score":<0-50>,"text_score":<0-30>,"tech_score":<0-20>,"total_score":<0-100>,"analysis":"概括","text_feedback":"文字评价","image_feedback":"截图评价"}}"""
    return prompt


def build_text_only_prompt(text):
    prompt = f"""你是严格的大学实验报告评分助手。以下报告没有提交截图。

## 学生文字内容
```
{text[:4000]}
```

## 定等级
A级(85-90): 目的+步骤+总结齐全，内容充实有深度，技术细节丰富
B级(70-79): 基本完整，有一定深度
C级(60-69): 包含主要要素但简略或有重复
D级(40-59): 内容残缺或极其简略
E级(0-39): 几乎无内容

## 硬扣分规则
- 无截图：总分上限75
- 文字<300字：总分上限30
- 文字300-500字：总分上限50
- 无总结：总分上限60
- 大量重复：总分上限50

## 输出要求
纯JSON：
{{"grade":"A/B/C/D/E","content_score":<0-30>,"tech_score":<0-25>,"summary_score":<0-25>,"format_score":<0-20>,"total_score":<0-100>,"analysis":"概括","text_feedback":"具体评价"}}"""
    return prompt


def is_zhipu_available():
    """检查智谱AI API Key是否配置"""
    return bool(ZHIPU_API_KEY)


def get_available_vision_models():
    """返回可用模型信息"""
    if ZHIPU_API_KEY:
        return [ZHIPU_MODEL]
    return []


def image_to_base64(image_path):
    """将图片转为base64编码"""
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def analyze_image_with_zhipu(image_path, text):
    """调用智谱AI视觉模型分析单张截图（结合文字内容）"""
    if not ZHIPU_API_KEY:
        return {"error": "未配置智谱API Key"}

    try:
        image_b64 = image_to_base64(image_path)
        image_mime = "image/png"
        if image_path.lower().endswith(('.jpg', '.jpeg')):
            image_mime = "image/jpeg"
        elif image_path.lower().endswith('.gif'):
            image_mime = "image/gif"
        elif image_path.lower().endswith('.webp'):
            image_mime = "image/webp"

        data_url = f"data:{image_mime};base64,{image_b64}"
        grading_prompt = build_grading_prompt(text)

        payload = {
            "model": ZHIPU_MODEL,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": data_url}},
                        {"type": "text", "text": grading_prompt}
                    ]
                }
            ],
            "temperature": 0.1,
            "max_tokens": 1024
        }

        headers = {
            "Authorization": f"Bearer {ZHIPU_API_KEY}",
            "Content-Type": "application/json"
        }

        resp = requests.post(
            f"{ZHIPU_BASE_URL}chat/completions",
            headers=headers,
            json=payload,
            timeout=120
        )

        if resp.status_code != 200:
            return {"error": f"智谱API返回错误 ({resp.status_code}): {resp.text[:200]}"}

        result = resp.json()
        content = result.get("choices", [{}])[0].get("message", {}).get("content", "")

        if not content:
            return {"error": "智谱API返回为空"}

        # 从回复中提取JSON（支持纯JSON或markdown代码块包裹）
        cleaned = re.sub(r'^```(?:json)?\s*|\s*```$', '', content.strip(), flags=re.MULTILINE)
        json_match = re.search(r'\{.*"total_score".*\}', cleaned, re.DOTALL)
        if json_match:
            try:
                data = json.loads(json_match.group())
                return data
            except json.JSONDecodeError as e:
                return {"error": f"JSON解析失败: {str(e)}, 原始输出: {content[:200]}"}
        else:
            return {"error": f"无法解析模型输出: {content[:300]}"}

    except requests.Timeout:
        return {"error": "智谱API请求超时（120s），图片可能过大"}
    except Exception as e:
        return {"error": f"分析失败: {str(e)}"}


def build_text_only_prompt(text):
    """构建纯文字评分提示词（无截图时使用，严格压低分数）"""
    text_excerpt = text[:4000] if text else "(无文字内容)"

    prompt = f"""你是一个严格的大学实验报告评分助手。以下实验报告完全没有提交截图，请根据纯文字内容评分。

## 学生文字内容
```
{text_excerpt}
```

## 评分标准（满分100分）

### 内容完整性 (0-30分)
- 结构完整，目的/步骤/总结齐全，内容充实: 25-30分
- 包含大部分要素但不够完整: 15-24分
- 内容残缺或过于简略: 0-14分

### 技术深度 (0-25分)
- 有具体命令、配置、代码片段: 20-25分
- 有技术描述但不够具体: 10-19分
- 几乎没有技术内容: 0-9分

### 总结反思 (0-25分)
- 总结有深度，问题分析到位，有解决方案: 20-25分
- 有总结但比较表面: 10-19分
- 没有总结或只有套话: 0-9分

### 格式与规范 (0-20分)
- 内容精炼，无重复，排版规范: 15-20分
- 少量重复，基本规范: 8-14分
- 大量重复内容或排版混乱: 0-7分

## 严格扣分规则
- 文字<300字：总分上限35分
- 文字300-500字：总分上限50分
- 完全没有实验总结：总分上限60分
- 内容大量重复（占全文30%以上）：总分上限45分
- 没有提交任何截图：直接扣15分，且总分上限70分

## 输出要求
纯JSON格式：
{{"content_score":<0-30>,"tech_score":<0-25>,"summary_score":<0-25>,"format_score":<0-20>,"total_score":<0-100>,"analysis":"优缺点概括","text_feedback":"具体评价"}}"""
    return prompt


def ai_grade_text_only(text):
    """无截图时，使用AI纯文字评分"""
    if not ZHIPU_API_KEY:
        return auto_grade_fallback(text, [])

    try:
        prompt = build_text_only_prompt(text)

        payload = {
            "model": ZHIPU_MODEL,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt}
                    ]
                }
            ],
            "temperature": 0.1,
            "max_tokens": 1024
        }

        headers = {
            "Authorization": f"Bearer {ZHIPU_API_KEY}",
            "Content-Type": "application/json"
        }

        print("  📝 正在AI文字评分（无截图）...")
        resp = requests.post(
            f"{ZHIPU_BASE_URL}chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )

        if resp.status_code != 200:
            print(f"  ⚠️ AI文字评分API错误: {resp.status_code}")
            return auto_grade_fallback(text, [])

        result = resp.json()
        content = result.get("choices", [{}])[0].get("message", {}).get("content", "")

        if not content:
            return auto_grade_fallback(text, [])

        cleaned = re.sub(r'^```(?:json)?\s*|\s*```$', '', content.strip(), flags=re.MULTILINE)
        json_match = re.search(r'\{.*"total_score".*\}', cleaned, re.DOTALL)
        if json_match:
            data = json.loads(json_match.group())
            score = data.get("total_score", 50)
            analysis = data.get("analysis", "")
            feedback = data.get("text_feedback", "")

            comments = []
            if score >= 90:
                comments.append("🎉 文字内容质量优秀！")
            elif score >= 75:
                comments.append("✅ 文字内容良好。")
            elif score >= 60:
                comments.append("⚠️ 文字内容基本合格。")
            elif score >= 40:
                comments.append("⚠️ 文字内容不够充分，需要补充。")
            else:
                comments.append("❌ 文字内容严重不足。")

            comments.append(f"⚠️ 该报告未提交任何实验截图，已根据评分规则扣分")
            if analysis:
                comments.append(f"📝 {analysis}")
            if feedback:
                comments.append(f"📝 评价: {feedback}")

            print(f"  ✅ AI文字评分完成: {score}分")
            return score, "；".join(comments)

        return auto_grade_fallback(text, [])

    except Exception as e:
        print(f"  ⚠️ AI文字评分异常: {e}")
        return auto_grade_fallback(text, [])


def calibrate_score(score, text, images):
    '''AI评分后硬性校准，强制拉开区分度'''
    import re
    deductions = []
    
    # 截图数量惩罚
    img_count = len(images)
    if img_count == 0:
        score = min(score, 30)
        deductions.append("无截图: 上限30分")
    elif img_count == 1:
        score -= 15
        deductions.append("仅1张截图: -15分")
    elif img_count == 2:
        score -= 8
        deductions.append("仅2张截图: -8分")
    elif img_count >= 4:
        score += 5
        deductions.append("截图充足(>=4张): +5分")
    
    # 内容长度惩罚
    text_len = len(text)
    if text_len < 300:
        score = min(score, 35)
        deductions.append("文字过短(<300字): 上限35分")
    elif text_len < 500:
        score = min(score, 55)
        deductions.append("文字偏短(<500字): 上限55分")
    elif text_len < 800:
        score -= 5
        deductions.append("文字略少(<800字): -5分")
    elif text_len > 2000:
        score += 3
        deductions.append("内容充实(>2000字): +3分")
    
    # 实验总结检查
    if not re.search(r'实验总结|实验心得|实验小结|心得体会', text):
        score -= 10
        deductions.append("缺少实验总结: -10分")
    
    # 遇到问题检查
    if not re.search(r'遇到的问题|问题及解决|错误|报错|故障|问题', text):
        score -= 5
        deductions.append("未记录遇到的问题: -5分")
    
    # 命令/代码检查
    if not re.search(r'start-dfs|start-yarn|jps|hadoop|hive|ssh|命令|配置', text):
        score -= 5
        deductions.append("缺少关键命令/代码: -5分")
    
    # 重复内容检查
    paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 20]
    if len(paragraphs) > 3:
        unique = set(paragraphs)
        if len(unique) < len(paragraphs) * 0.6:
            score -= 15
            deductions.append("内容大量重复: -15分")
    
    score = max(0, min(100, round(score)))
    return score, deductions


def ai_grade_with_zhipu(text, images):
    """
    使用智谱AI视觉模型进行图片+文字综合评分
    返回 (score, comment)
    """
    # ===== 没有截图时，走AI纯文字评分（严格） =====
    if not images:
        print("  ⚠️ 该文档没有截图，使用AI文字评分...")
        return ai_grade_text_only(text)

    # ===== 有截图时，每张截图带文字一起综合评分 =====
    all_findings = []
    total_scores = []
    text_feedbacks = []
    image_feedbacks = []
    any_error = False

    for idx, img in enumerate(images):
        image_path = img["path"]
        print(f"  📤 正在综合评分 截图{idx+1}/{len(images)}: {img['filename']}")
        result = analyze_image_with_zhipu(image_path, text)

        if "error" in result:
            print(f"  ⚠️ 截图 {idx+1} 分析失败: {result['error']}")
            any_error = True
            continue

        score = result.get("total_score", 60)
        total_scores.append(score)

        # AI返回的字段确保转为字符串，防止dict等不可hash类型导致报错
        raw_text_fb = result.get("text_feedback", "")
        raw_image_fb = result.get("image_feedback", "")
        text_fb = str(raw_text_fb) if raw_text_fb else ""
        image_fb = str(raw_image_fb) if raw_image_fb else ""
        if text_fb:
            text_feedbacks.append(text_fb)
        if image_fb:
            image_feedbacks.append(image_fb)

        all_findings.append({
            "index": idx + 1,
            "filename": img["filename"],
            "score": score,
            "image_score": result.get("image_score", 0),
            "text_score": result.get("text_score", 0),
            "tech_score": result.get("tech_score", 0),
            "analysis": result.get("analysis", ""),
            "text_feedback": text_fb,
            "image_feedback": image_fb,
        })

        print(f"  ✅ 截图{idx+1}: 总分{score} (截图{result.get('image_score',0)}分+文字{result.get('text_score',0)}分)")

    # 综合评分
    if total_scores:
        avg_score = sum(total_scores) / len(total_scores)
        final_score = round(max(0, min(100, avg_score)))
    else:
        print("  ⚠️ 所有截图分析失败，回退到文字评分")
        return ai_grade_text_only(text)

    # 计算平均分项
    text_score_avg = 0
    image_score_avg = 0
    tech_score_avg = 0
    if all_findings:
        text_score_avg = sum(f.get("text_score", 0) for f in all_findings) / len(all_findings)
        image_score_avg = sum(f.get("image_score", 0) for f in all_findings) / len(all_findings)
        tech_score_avg = sum(f.get("tech_score", 0) for f in all_findings) / len(all_findings)

    # 硬性校准 - 强制拉开区分度
    deductions = []
    old_score = final_score
    final_score, deductions = calibrate_score(final_score, text, images)
    if deductions:
        print(f"  🔧 校准: {old_score} -> {final_score} ({'; '.join(deductions)})")

    comments = []
    if final_score >= 90:
        comments.append("🎉 实验完成度极高！")
    elif final_score >= 75:
        comments.append("✅ 实验基本完成。")
    elif final_score >= 60:
        comments.append("⚠️ 实验部分完成，有改进空间。")
    else:
        comments.append("❌ 实验存在明显问题，需要完善。")

    comments.append(f"截图 {image_score_avg:.0f}/50 + 内容 {text_score_avg:.0f}/30 + 技术 {tech_score_avg:.0f}/20 = {old_score} 分(校准后{final_score}分)")

    if text_feedbacks:
        seen = set()
        for fb in text_feedbacks:
            if fb not in seen:
                seen.add(fb)
                if len(seen) <= 2:
                    comments.append(f"📝 文字: {fb}")

    if image_feedbacks:
        seen = set()
        for fb in image_feedbacks:
            if fb not in seen:
                seen.add(fb)
                if len(seen) <= 2:
                    comments.append(f"🖼️ 截图: {fb}")

    if len(all_findings) > 1:
        for f in all_findings:
            analysis = f.get("analysis", "")
            if analysis:
                comments.append(f"截图{f['index']}: {analysis}")

    return final_score, "；".join(comments)


def auto_grade_fallback(text, images):
    """
    自动评分（回退方案）：基于文本关键词匹配 + 图片数量统计
    当API完全不可用时使用。评分从严。
    """
    has_images = len(images) > 0
    score = 30 if has_images else 10  # 没有截图基础分极低

    has_summary = bool(re.search(r'实验总结|实验心得|实验小结', text))
    if has_summary:
        score += 10

    has_problems = bool(re.search(r'遇到的问题|问题及解决|错误|报错|故障', text))
    if has_problems:
        score += 10

    if len(images) >= 3:
        score += 20
    elif len(images) >= 1:
        score += 10

    has_code = bool(re.search(r'```|start-dfs|start-yarn|jps|hadoop|hive', text))
    if has_code:
        score += 5

    word_count = len(text)
    if word_count > 2000:
        score += 5
    elif word_count > 1000:
        score += 3

    has_purpose = bool(re.search(r'实验目的|实验目标', text))
    has_steps = bool(re.search(r'实验步骤|操作步骤|实验原理', text))
    if has_purpose and has_steps:
        score += 5

    if word_count < 500:
        score -= 15
    elif word_count < 300:
        score -= 30

    # 没有截图直接上限扣分
    if not has_images:
        score = min(score, 50)  # 无截图最高50分

    score = max(0, min(100, score))

    comments = []
    if score >= 90:
        comments.append("[⚠️ 关键词评分] 实验完成度极高，内容完整。")
    elif score >= 75:
        comments.append("[⚠️ 关键词评分] 实验基本完成，内容较为完整。")
    elif score >= 60:
        comments.append("[⚠️ 关键词评分] 实验部分完成，有改进空间。")
    elif score >= 40:
        comments.append("[⚠️ 关键词评分] 实验内容不足，需要补充完善。")
    else:
        comments.append("[⚠️ 关键词评分] 实验严重不足，请重新提交。")

    if not has_images:
        comments.append("❌ 缺少实验截图，这是严重扣分项！")
    if has_summary:
        comments.append("有实验总结。")
    if has_problems:
        comments.append("记录了遇到的问题。")
    if not has_code:
        comments.append("建议补充关键命令或代码。")

    return score, "；".join(comments)


# ========== 文档解析模块 ==========

def extract_student_info(doc, filename=""):
    """从Word文档中提取学生信息（支持多种文档格式）"""
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text.strip() + " "

    info = {
        "student_id": "",
        "name": "",
        "class_name": "",
        "course": "",
        "experiment": "",
    }

    # 学号：多种格式
    id_patterns = [
        r'学\s*号[：:\s/]*\s*_*(\d{10,12})',           # 学号：2024... / 学号____2024...
        r'学\s*号[：:\s]*(\d{10,12})',                    # 学号：2024...
        r'学号/班级[：:\s]*_*(\d{10,12})',                # 学号/班级：____2024...
        r'学号[：:\s]*_*(\d{10,12})',                      # 学号：____2024...
    ]
    for p in id_patterns:
        m = re.search(p, text)
        if m:
            info["student_id"] = m.group(1).strip()
            break

    # 如果文档内没找到，从文件名提取（文件名常有学号）
    if not info["student_id"] and filename:
        file_match = re.search(r'(\d{10,12})', filename)
        if file_match:
            info["student_id"] = file_match.group(1)

    # 姓名：多种格式
    name_patterns = [
        r'姓\s*名[：:\s]*_*([一-龥]{2,4})_*',             # 姓名：张三____
        r'姓\s*名\s+([一-龥]{2,4})',                       # 姓   名   张三
        r'实验人[：:\s]*_*([一-龥]{2,4})_*',               # 实验人：____张三____
        r'姓名[：:\s]*([一-龥]{2,4})',                      # 姓名：张三
    ]
    for p in name_patterns:
        m = re.search(p, text)
        if m:
            info["name"] = m.group(1).strip()
            break

    # 也从文件名提取姓名（文件名常有姓名）
    if not info["name"] and filename:
        # 文件名在学号后面的可能是姓名，如 "20233822012胡小虎xxx.docx"
        name_in_file = re.search(r'\d{10,12}([一-龥]{2,4})', filename)
        if name_in_file:
            info["name"] = name_in_file.group(1)

    # 班级：多种格式
    class_patterns = [
        r'学号/班级[：:\s]*_*\d{10,12}_*[/_]*([^\s_]{3,10})',   # 学号/班级：____2024.../23物联网
        r'班\s*级[：:\s]*_*([^\s]{2,10})',
        r'班级[：:\s]*_*([^\s_]{2,10})',
    ]
    for p in class_patterns:
        m = re.search(p, text)
        if m:
            info["class_name"] = m.group(1).strip()
            break

    # 课程名称
    course_patterns = [
        r'课程名称[：:\s]+([^\s]{2,20})',
        r'课程[：:\s]*([^\s]{2,20})',
    ]
    for p in course_patterns:
        m = re.search(p, text)
        if m:
            info["course"] = m.group(1).strip()
            break

    # 实验名称（优先文档标题，再搜标签）
    first_para = ""
    for para in doc.paragraphs:
        if para.text.strip():
            first_para = para.text.strip()
            break

    exp_patterns = [
        r'实验[一二三四五六七八九十]+[：:]\s*(.+?)(?:\n|$)',
        r'实验名称[：:\s]+([^\s]{2,30})',
        r'实验[：:\s]*([^\s]{2,30})',
    ]
    for p in exp_patterns:
        m = re.search(p, text)
        if m:
            info["experiment"] = m.group(1).strip()
            break

    # 如果正则没找到，用文档第一段标题作为实验名称
    if not info["experiment"] and first_para and len(first_para) > 4:
        info["experiment"] = first_para

    return info, text


def extract_images(doc, doc_id):
    """从Word文档中提取图片，保存到images目录"""
    images = []
    for rel_id, rel in doc.part.rels.items():
        if "image" in str(rel.reltype):
            image_data = rel.target_part.blob
            ext = os.path.splitext(rel.target_ref)[1] or ".png"
            fname = f"{doc_id}_{rel_id}{ext}"
            fpath = os.path.join(IMAGE_FOLDER, fname)
            with open(fpath, "wb") as f:
                f.write(image_data)
            try:
                img_url = url_for('serve_image', filename=fname)
            except RuntimeError:
                img_url = f"/images/{fname}"
            images.append({
                "filename": fname,
                "path": fpath,
                "size": len(image_data),
                "url": img_url
            })
    return images


# ========== Excel 导出模块 ==========

def export_to_excel(results, filename="成绩汇总表.xlsx"):
    """将评分结果导出为Excel文件"""
    wb = Workbook()
    ws = wb.active
    ws.title = "成绩汇总"

    # 表头样式
    header_font = Font(name='微软雅黑', bold=True, color='FFFFFF', size=12)
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_alignment = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 写入表头
    headers = ['序号', '学号', '姓名', '班级', '课程', '实验名称', '分数', '评语', '截图数量', '批改时间']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # 写入数据
    for i, result in enumerate(results, 1):
        row = i + 1
        ws.cell(row=row, column=1, value=i).border = thin_border
        ws.cell(row=row, column=2, value=result.get('student_id', '')).border = thin_border
        ws.cell(row=row, column=3, value=result.get('name', '')).border = thin_border
        ws.cell(row=row, column=4, value=result.get('class_name', '')).border = thin_border
        ws.cell(row=row, column=5, value=result.get('course', '')).border = thin_border
        ws.cell(row=row, column=6, value=result.get('experiment', '')).border = thin_border
        ws.cell(row=row, column=7, value=result.get('score', 0)).border = thin_border
        ws.cell(row=row, column=8, value=result.get('comment', '')).border = thin_border
        ws.cell(row=row, column=9, value=result.get('image_count', 0)).border = thin_border
        ws.cell(row=row, column=10, value=datetime.now().strftime('%Y-%m-%d %H:%M')).border = thin_border

        # 成绩列特殊颜色
        score_cell = ws.cell(row=row, column=7)
        score = result.get('score', 0)
        if score >= 90:
            score_cell.fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
        elif score >= 75:
            score_cell.fill = PatternFill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')
        elif score >= 60:
            score_cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
        else:
            score_cell.fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')

    # 设置列宽
    col_widths = [6, 16, 10, 16, 16, 30, 8, 40, 10, 18]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w

    filepath = os.path.join(BASE_DIR, filename)
    wb.save(filepath)
    return filepath


# ========== Flask 路由 ==========

@app.route('/')
def index():
    """首页：显示已上传的文档列表"""
    documents = []
    if os.path.exists(REPORTS_FOLDER):
        for fname in os.listdir(REPORTS_FOLDER):
            if fname.endswith('.docx'):
                fpath = os.path.join(REPORTS_FOLDER, fname)
                mtime = datetime.fromtimestamp(os.path.getmtime(fpath))
                documents.append({
                    'filename': fname,
                    'path': fpath,
                    'modified': mtime.strftime('%Y-%m-%d %H:%M'),
                    'size': f"{os.path.getsize(fpath) / 1024:.1f} KB"
                })

    all_results = load_results()
    return render_template('index.html',
                           documents=documents,
                           results=all_results)


@app.route('/upload', methods=['POST'])
def upload():
    """上传Word文档"""
    if 'files' not in request.files:
        flash('请选择文件', 'error')
        return redirect(url_for('index'))

    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        flash('请选择文件', 'error')
        return redirect(url_for('index'))

    uploaded = []
    for f in files:
        if not f.filename.endswith('.docx'):
            continue
        # 处理文件夹上传带来的相对路径（取最后一部分作为文件名）
        original_name = f.filename
        if '/' in original_name:
            original_name = original_name.rsplit('/', 1)[-1]
        elif '\\' in original_name:
            original_name = original_name.rsplit('\\', 1)[-1]

        fpath = os.path.join(REPORTS_FOLDER, original_name)
        # 同名文件加序号
        counter = 1
        while os.path.exists(fpath):
            name, ext = os.path.splitext(original_name)
            fpath = os.path.join(REPORTS_FOLDER, f"{name}_{counter}{ext}")
            counter += 1

        f.save(fpath)
        uploaded.append(os.path.basename(fpath))

    if uploaded:
        flash(f'成功上传 {len(uploaded)} 个文档：{", ".join(uploaded)}', 'success')
    else:
        flash('未上传任何 .docx 文件', 'error')

    return redirect(url_for('index'))


@app.route('/grade')
def grade():
    """评分页面：解析文档并展示评分结果"""
    # 检查智谱AI API状态
    api_available = is_zhipu_available()
    vision_models = get_available_vision_models()
    grading_mode = "智谱AI视觉评分" if api_available else "关键词回退评分"

    if not api_available:
        flash('⚠️ 未配置智谱API Key，使用关键词回退评分。请在环境变量中设置 ZHIPU_API_KEY', 'warning')

    all_results = []

    for fname in os.listdir(REPORTS_FOLDER):
        if not fname.endswith('.docx'):
            continue

        fpath = os.path.join(REPORTS_FOLDER, fname)
        doc_id = os.path.splitext(fname)[0]

        try:
            doc = docx.Document(fpath)
            info, text = extract_student_info(doc, fname)
            images = extract_images(doc, doc_id)

            # 优先使用AI视觉评分
            if api_available and images:
                print(f"\n📝 正在AI评分: {fname} ({len(images)}张截图)")
                score, comment = ai_grade_with_zhipu(text, images)
            else:
                score, comment = auto_grade_fallback(text, images)

            result = {
                'id': doc_id,
                'filename': fname,
                **info,
                'score': score,
                'comment': comment,
                'images': images,
                'image_count': len(images),
                'text_preview': text[:500],
            }
            all_results.append(result)
        except Exception as e:
            all_results.append({
                'id': doc_id,
                'filename': fname,
                'student_id': '',
                'name': f'[解析失败: {str(e)}]',
                'class_name': '',
                'course': '',
                'experiment': '',
                'score': 0,
                'comment': f'文档解析失败：{str(e)}',
                'images': [],
                'image_count': 0,
                'text_preview': '',
            })

    print(f"\n📊 评分完成: {len(all_results)} 个文档")
    save_results(all_results)

    return render_template('grade.html',
                           results=all_results,
                           grading_mode=grading_mode,
                           api_available=api_available,
                           api_model=ZHIPU_MODEL,
                           vision_models=vision_models)


@app.route('/update_score', methods=['POST'])
def update_score():
    """手动调整分数"""
    data = request.get_json()
    doc_id = data.get('id')
    new_score = data.get('score')
    new_comment = data.get('comment')

    results = load_results()
    for r in results:
        if r['id'] == doc_id:
            r['score'] = int(new_score)
            if new_comment:
                r['comment'] = new_comment
            break

    save_results(results)
    return jsonify({'status': 'ok'})


@app.route('/export')
def export():
    """导出为Excel"""
    results = load_results()
    if not results:
        flash('没有评分数据可以导出', 'error')
        return redirect(url_for('index'))

    filepath = export_to_excel(results)
    filename = os.path.basename(filepath)
    flash(f'Excel已导出：{filename}', 'success')

    return send_from_directory(BASE_DIR, filename, as_attachment=True)


@app.route('/export_html')
def export_html():
    """导出为HTML报告"""
    results = load_results()
    if not results:
        flash('没有评分数据可以导出', 'error')
        return redirect(url_for('index'))

    html = render_template('report.html', results=results, now=datetime.now())
    report_path = os.path.join(BASE_DIR, '批改报告.html')
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)

    flash('HTML报告已生成：批改报告.html', 'success')
    return redirect(url_for('index'))


@app.route('/images/<filename>')
def serve_image(filename):
    """提供图片访问"""
    return send_from_directory(IMAGE_FOLDER, filename)


@app.route('/clear', methods=['POST'])
def clear_all():
    """清空所有数据和文件"""
    for folder in [REPORTS_FOLDER, IMAGE_FOLDER, UPLOAD_FOLDER]:
        for f in os.listdir(folder):
            fpath = os.path.join(folder, f)
            try:
                if os.path.isfile(fpath):
                    os.unlink(fpath)
            except Exception:
                pass
    results_file = os.path.join(BASE_DIR, '.grading_results.json')
    if os.path.exists(results_file):
        os.unlink(results_file)

    flash('已清空所有数据', 'success')
    return redirect(url_for('index'))


@app.route('/api/api_status')
def api_status_api():
    """API: 返回智谱AI连接状态"""
    available = is_zhipu_available()
    models = get_available_vision_models()
    return jsonify({
        "available": available,
        "provider": "zhipuai",
        "model": ZHIPU_MODEL,
        "vision_models": models
    })


# ========== 数据持久化 ==========

def save_results(results):
    """保存评分结果到JSON"""
    data = []
    for r in results:
        d = {k: v for k, v in r.items() if k != 'images'}
        d['image_count'] = len(r.get('images', []))
        d['text_preview'] = r.get('text_preview', '')[:500]
        data.append(d)

    filepath = os.path.join(BASE_DIR, '.grading_results.json')
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_results():
    """从JSON加载评分结果"""
    filepath = os.path.join(BASE_DIR, '.grading_results.json')
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []


# ========== 启动 ==========

def open_browser():
    """延迟打开浏览器"""
    time.sleep(2)
    try:
        webbrowser.open('http://localhost:8900')
    except Exception:
        pass

if __name__ == '__main__':
    is_packaged = getattr(sys, 'frozen', False)

    print("=" * 60)
    print("  📚 作业批改系统")
    print(f"  当前时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"  待批改目录: {REPORTS_FOLDER}")
    print(f"  图片目录: {IMAGE_FOLDER}")
    if ZHIPU_API_KEY:
        print(f"  🤖 评分引擎: 智谱AI ({ZHIPU_MODEL})")
    else:
        print("  ⚠️ 未配置API Key")
    print("=" * 60)
    print("  ✅ 启动成功！请在浏览器中打开：")
    print("  👉  http://localhost:8900")
    print("=" * 60)
    print("  关闭此窗口即关闭服务")
    print()

    # 自动打开浏览器
    threading.Thread(target=open_browser, daemon=True).start()

    try:
        app.run(debug=not is_packaged, host='127.0.0.1', port=8900)
    except KeyboardInterrupt:
        pass
    except Exception as e:
        print(f"\n❌ 启动失败: {e}")
        input("\n按回车键退出...")
